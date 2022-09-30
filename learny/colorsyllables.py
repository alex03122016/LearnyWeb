#! python3
# -*- coding: utf-8 -*-
# aim: help students learning to read by coloring red every second syllable
# idea was seen in non free educational program milbenberger silbengenerator
# what the program does:
# reads an existing docx file with tables
# in words with more than one syllaba colors every second syllaba red
# creates a docx file to print the result
# TODO: create a document that keeps original formatting on run level i.e. with pictures and special formatting but with colored syllabas

def color_syllables(inputtext):
	import subprocess
	import os
	try:
		import docxprint
	except ImportError:
		from learny import docxprint
	#python-docx module code: pip install python-docx
	import docx
	from docx.shared import RGBColor

	#code: pip install PyHyphen
	from hyphen import Hyphenator
	de_DE = Hyphenator('de_DE')

	print('color syllables')
	print('input Text: ' + inputtext)


	# variables and lists used
	#save_path = os.path.join(os.path.expanduser('~'),'python-project' ,'kivy-test', 'learny', __name__ + 'fileTitle.docx')

	i = 0
	word_print = []
	Aufgabe = {
		"Kopfzeile": "Name: 				Klasse: 				Datum:  \n ",
		"Titel": "",
		"1. Aufgabe": "Lies den Text! Tippe bei jeder Silbe auf den Tisch!\n",
		"Hinweise": "Hier sind die Wörter aus den Lücken: \n",
		"Rätselwörter": "Hier ein paar Rätselwörter aus dem Text: \n",
	}
	# create document object
	doc = docx.Document()
	save_path = docxprint.docx_print(Doc= doc, save= 'colorsyllables')
	#get the input text
	text = inputtext
	#only work on text with '' specific words
	if text.find('') != -1:
		hyph_print = []
		word_print =[]
		hyph =[]

	#seperate text into words and put them in list hyph_input
		hyph_input = text.split()
		docxprint.docx_print(printText=Aufgabe["Kopfzeile"], Doc=doc)
		docxprint.docx_print(printText=Aufgabe["1. Aufgabe"], Bold=True, Doc=doc)
		paragraph = docxprint.docx_print(Doc=doc)

	#work on separated words
		for w in hyph_input:
			hyph_print = []
			i = 1
			hyph = de_DE.syllables(w)
			if hyph == []:
				#paragraph.add_run(w+' ')
				docxprint.docx_print(printText=w + ' ', Paragraph=paragraph, Doc=doc)	#adds the text
			for syl in hyph:
				if i%2 == 0:
					docxprint.docx_print(printText=syl, color="red", Paragraph=paragraph, Doc=doc)
				else :
					#paragraph.add_run(syl)
					docxprint.docx_print(printText=syl, Paragraph=paragraph, Doc=doc)	#adds the text
				i +=1
			docxprint.docx_print(printText=' ', Paragraph=paragraph, Doc=doc)	#adds the text

	#save the result in absolute path
	doc.save(save_path)

	return

input= """
In vielen Getränken in Deutschland ist zu viel Zucker. Das hat die Organisation „Foodwatch“ festgestellt. Wenn man zu viel Zucker isst, kann man zu dick und krank werden.

Die Organisation „Foodwatch“ hat 463 Limonaden, Energy Drinks, Saft-Schorlen und Eis-Tees untersucht. „Foodwatch“ ist ein englisches Wort und heißt Lebens-Mittel-Überwachung. Nur sechs Getränke waren nicht gezuckert. Das süßeste Getränk – ein Energy-Drink – hatte in einer halben Liter Dose 78 Gramm Zucker. Das sind 26 Stück Würfel-Zucker.
Die Welt-Gesundheits-Organisation hat empfohlen: Ein Mensch sollte nicht mehr als sechs Tee-Löffel Zucker am Tag essen. Wenn man mehr Zucker isst, können die Zähne krank werden oder man wird zu dick. Eine große Gefahr ist auch die Zucker-Krankheit. Sie wird Diabetes genannt. Daran kann der Mensch sterben.
In dem Land Groß-Britannien soll es deshalb ab 2018 eine Zucker-Steuer geben, damit die Hersteller weniger süße Getränke verkaufen.
„Foodwatch“ fordert: Auch in Deutschland soll es eine Zucker-Steuer geben. Die Bundes-Regierung lehnt das ab. Der Minister für Ernährung, Christian Schmidt, sagt: Straf-Steuern auf Lebens-Mittel sind der falsche Weg. Die Menschen sollen besser über gesunde Lebens-Mittel informiert werden. Am besten schon in der Schule.

"""

color_syllables(inputtext= input)
